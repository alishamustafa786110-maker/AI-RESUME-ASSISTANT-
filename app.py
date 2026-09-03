"""
ATS Resume Analyzer
--------------------
A Streamlit app that lets a user upload a resume (PDF / DOCX / TXT),
optionally paste a target job description, and get back:
  - An ATS compatibility score (0-100)
  - A breakdown of the score
  - Strengths
  - Concrete improvement suggestions
  - Missing / weak keywords vs. the job description
  - Formatting issues that can trip up ATS parsers

Powered by Google's Gemini API (model: gemini-3.5-flash) using the
`google-genai` SDK's Interactions API with structured JSON output.
"""

import io
import json
import os

import streamlit as st
from pydantic import BaseModel, Field
from pypdf import PdfReader
from docx import Document
from google import genai
from google.genai import errors as genai_errors


# --------------------------------------------------------------------------
# Config
# --------------------------------------------------------------------------
MODEL_NAME = "gemini-3.5-flash"

st.set_page_config(
    page_title="ATS Resume Analyzer",
    page_icon="📄",
    layout="centered",
)


# --------------------------------------------------------------------------
# Structured output schema
# --------------------------------------------------------------------------
class ScoreBreakdown(BaseModel):
    keyword_match: int = Field(description="0-100 score for how well the resume's keywords match the target role/job description.")
    formatting: int = Field(description="0-100 score for ATS-friendly formatting (no tables/columns issues, standard fonts, parseable structure).")
    structure: int = Field(description="0-100 score for section structure and completeness (contact info, experience, education, skills, etc).")
    clarity: int = Field(description="0-100 score for clarity and impact of bullet points (quantified achievements, action verbs).")


class ATSAnalysis(BaseModel):
    ats_score: int = Field(description="Overall ATS compatibility score from 0 to 100.")
    score_breakdown: ScoreBreakdown
    summary: str = Field(description="A 2-3 sentence overall summary of the resume's ATS readiness.")
    strengths: list[str] = Field(description="3-6 specific things the resume does well.")
    improvements: list[str] = Field(description="5-10 specific, actionable improvements, ordered by impact.")
    missing_keywords: list[str] = Field(description="Important keywords/skills missing or underrepresented, especially relative to the job description if provided. Empty list if none.")
    formatting_issues: list[str] = Field(description="Specific formatting problems that could confuse an ATS parser (e.g. tables, images, headers/footers, non-standard section titles). Empty list if none found.")


# --------------------------------------------------------------------------
# Resume text extraction
# --------------------------------------------------------------------------
def extract_text_from_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages_text = []
    for page in reader.pages:
        text = page.extract_text() or ""
        pages_text.append(text)
    return "\n".join(pages_text).strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())
    return "\n".join(parts).strip()


def extract_text_from_txt(file_bytes: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return ""


def extract_resume_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    file_bytes = uploaded_file.getvalue()

    if name.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif name.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif name.endswith(".txt"):
        return extract_text_from_txt(file_bytes)
    else:
        raise ValueError("Unsupported file type. Please upload a PDF, DOCX, or TXT file.")


# --------------------------------------------------------------------------
# Gemini call
# --------------------------------------------------------------------------
def get_client(api_key: str) -> genai.Client:
    return genai.Client(api_key=api_key)


def build_prompt(resume_text: str, job_description: str) -> str:
    prompt = f"""You are an expert ATS (Applicant Tracking System) resume reviewer and career coach.

Analyze the resume text below as an ATS parser would, and also from a human recruiter's
perspective. Be specific and actionable — reference actual phrases or sections from the
resume where relevant instead of generic advice.

RESUME TEXT:
\"\"\"
{resume_text}
\"\"\"
"""
    if job_description.strip():
        prompt += f"""
TARGET JOB DESCRIPTION (score keyword match and identify missing keywords against this):
\"\"\"
{job_description.strip()}
\"\"\"
"""
    else:
        prompt += """
No specific job description was provided. Evaluate keyword strength and relevance
generally for the apparent target role/industry based on the resume content itself.
"""
    prompt += """
Return your analysis strictly following the provided JSON schema.
"""
    return prompt


def analyze_resume(client: genai.Client, resume_text: str, job_description: str) -> ATSAnalysis:
    prompt = build_prompt(resume_text, job_description)

    interaction = client.interactions.create(
        model=MODEL_NAME,
        input=prompt,
        response_format={
            "type": "text",
            "mime_type": "application/json",
            "schema": ATSAnalysis.model_json_schema(),
        },
    )
    return ATSAnalysis.model_validate_json(interaction.output_text)


# --------------------------------------------------------------------------
# UI
# --------------------------------------------------------------------------
def render_score_gauge(score: int):
    if score >= 80:
        color = "🟢"
    elif score >= 60:
        color = "🟡"
    else:
        color = "🔴"
    st.metric(label=f"{color} ATS Score", value=f"{score}/100")
    st.progress(min(max(score, 0), 100) / 100)


def main():
    st.title("📄 ATS Resume Analyzer")
    st.caption("Upload your resume to get an ATS compatibility score and concrete improvement tips, powered by Gemini 3.5 Flash.")

    with st.sidebar:
        st.header("Settings")
        env_key = os.environ.get("GEMINI_API_KEY", "")
        secrets_key = ""
        try:
            secrets_key = st.secrets.get("GEMINI_API_KEY", "")
        except Exception:
            secrets_key = ""

        default_key = secrets_key or env_key
        api_key = st.text_input(
            "Gemini API Key",
            value=default_key,
            type="password",
            help="Get a free key at https://aistudio.google.com/apikey. "
                 "You can also set it as the GEMINI_API_KEY environment variable "
                 "or in Streamlit secrets so you don't have to paste it every time.",
        )
        st.markdown("---")
        st.caption(f"Model: `{MODEL_NAME}`")

    uploaded_file = st.file_uploader(
        "Upload your resume",
        type=["pdf", "docx", "txt"],
        help="PDF, DOCX, or TXT files are supported.",
    )

    job_description = st.text_area(
        "Target job description (optional, but recommended)",
        height=180,
        placeholder="Paste the job posting here to get a keyword-match score against this specific role...",
    )

    analyze_clicked = st.button("Analyze Resume", type="primary", disabled=uploaded_file is None)

    if analyze_clicked:
        if not api_key:
            st.error("Please enter your Gemini API key in the sidebar first.")
            st.stop()

        with st.spinner("Extracting text from your resume..."):
            try:
                resume_text = extract_resume_text(uploaded_file)
            except Exception as e:
                st.error(f"Couldn't read that file: {e}")
                st.stop()

        if not resume_text or len(resume_text.strip()) < 30:
            st.error(
                "Couldn't extract meaningful text from this file. "
                "If it's a scanned/image-based PDF, try exporting a text-based "
                "version from Word/Google Docs instead."
            )
            st.stop()

        with st.spinner("Analyzing your resume with Gemini... this can take a few seconds."):
            try:
                client = get_client(api_key)
                analysis = analyze_resume(client, resume_text, job_description)
            except genai_errors.ClientError as e:
                st.error(f"Gemini API rejected the request (check your API key and quota): {e}")
                st.stop()
            except genai_errors.APIError as e:
                st.error(f"Gemini API error: {e}")
                st.stop()
            except json.JSONDecodeError:
                st.error("The model returned a response that couldn't be parsed. Please try again.")
                st.stop()
            except Exception as e:
                st.error(f"Something went wrong: {e}")
                st.stop()

        st.success("Analysis complete!")
        render_score_gauge(analysis.ats_score)

        st.subheader("Summary")
        st.write(analysis.summary)

        cols = st.columns(4)
        breakdown = analysis.score_breakdown
        cols[0].metric("Keyword Match", f"{breakdown.keyword_match}/100")
        cols[1].metric("Formatting", f"{breakdown.formatting}/100")
        cols[2].metric("Structure", f"{breakdown.structure}/100")
        cols[3].metric("Clarity", f"{breakdown.clarity}/100")

        st.subheader("✅ Strengths")
        for item in analysis.strengths:
            st.markdown(f"- {item}")

        st.subheader("🛠️ Suggested Improvements")
        for i, item in enumerate(analysis.improvements, 1):
            st.markdown(f"{i}. {item}")

        if analysis.missing_keywords:
            st.subheader("🔑 Missing / Weak Keywords")
            st.markdown(", ".join(f"`{kw}`" for kw in analysis.missing_keywords))

        if analysis.formatting_issues:
            st.subheader("⚠️ Formatting Issues")
            for item in analysis.formatting_issues:
                st.markdown(f"- {item}")

        with st.expander("View extracted resume text"):
            st.text(resume_text)


if __name__ == "__main__":
    main()
